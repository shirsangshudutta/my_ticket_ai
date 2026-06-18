"""
doc_generator.py
Pure writer — zero AI logic.
Takes AI-generated section content and writes it into a .docx
that mirrors the uploaded template's heading structure exactly.

Handles:
- PlantUML blocks (@startuml ... @enduml) → formatted code block + caption
- Table content (rows separated by \\n, columns by |) → Word table
- Bullet lists (lines starting with -) → Word list bullets
- Plain paragraphs → Word paragraphs
"""

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import re


# ══════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════

def _add_table_from_text(doc, text: str):
    """
    Parse pipe-separated table text and write as a Word table.
    Format: "Col1 | Col2 | Col3\\nVal1 | Val2 | Val3"
    """
    rows = [r.strip() for r in text.strip().split("\n") if "|" in r]
    if not rows:
        doc.add_paragraph(text)
        return

    cols  = [c.strip() for c in rows[0].split("|")]
    table = doc.add_table(rows=len(rows), cols=len(cols), style="Table Grid")

    for r_idx, row_text in enumerate(rows):
        cells = [c.strip() for c in row_text.split("|")]
        for c_idx, cell_val in enumerate(cells):
            if c_idx < len(cols):
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = cell_val
                # Bold the header row
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True


def _add_plantuml_block(doc, plantuml_text: str, caption: str = ""):
    """
    Write a PlantUML block as a styled code paragraph with a caption.
    Note: actual diagram rendering requires PlantUML server integration.
    The code block is preserved so it can be rendered externally.
    """
    # Caption above the block
    if caption:
        p = doc.add_paragraph(caption)
        p.runs[0].bold   = True
        p.runs[0].italic = True

    # Code block style
    para = doc.add_paragraph(style="No Spacing")
    run  = para.add_run(plantuml_text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

    # Light grey shading on the paragraph
    pPr  = para._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F4F8")
    pPr.append(shd)

    doc.add_paragraph(
        "ℹ Render this PlantUML block at: https://www.plantuml.com/plantuml/uml/",
        style="No Spacing"
    ).runs[0].font.size = Pt(8)


def _write_content(doc, content: str):
    """
    Smart content writer — detects content type and renders appropriately.
    """
    # Convert non-string content to string
    if isinstance(content, dict):
        content = "\n".join(f"{k}: {v}" for k, v in content.items())
    elif isinstance(content, list):
        content = "\n".join(str(item) for item in content)
    elif not isinstance(content, str):
        content = str(content)
    
    if not content or not content.strip():
        doc.add_paragraph("[No content generated]")
        return

    # ── PlantUML diagram blocks ────────────────────────────────────────────
    plantuml_match = re.search(r"(@startuml.*?@enduml)", content, re.DOTALL | re.IGNORECASE)
    if plantuml_match:
        before = content[:plantuml_match.start()].strip()
        after  = content[plantuml_match.end():].strip()
        if before:
            _write_content(doc, before)
        _add_plantuml_block(doc, plantuml_match.group(1))
        if after:
            _write_content(doc, after)
        return

    # ── Pipe-separated table ───────────────────────────────────────────────
    lines = [l.strip() for l in content.strip().split("\n") if l.strip()]
    pipe_lines = [l for l in lines if "|" in l]
    if len(pipe_lines) >= 2 and len(pipe_lines) >= len(lines) * 0.5:
        _add_table_from_text(doc, "\n".join(pipe_lines))
        return

    # ── Mixed content: bullets, numbered steps, plain paragraphs ──────────
    for line in lines:
        if re.match(r"^\d+\.\s", line):
            # Numbered step
            doc.add_paragraph(line, style="List Number")
        elif line.startswith(("-", "•", "*")):
            # Bullet
            clean = line.lstrip("-•* ").strip()
            doc.add_paragraph(clean, style="List Bullet")
        elif line.startswith("[") and line.endswith("]"):
            # Placeholder / note — italicised
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        else:
            doc.add_paragraph(line)


# ══════════════════════════════════════════════════════════════════════════
# MAIN ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════

def generate_doc(
    source: str,
    targets: list,
    frequency: str,
    section_content: dict,          # {section_title: ai_generated_content}
    template_sections: list[dict],  # [{"level": int, "title": str}]
    output_path: str,
    interface_id: str = "",
) -> str:
    """
    Write the LLD Interface Architecture Document.

    Structure  = uploaded template headings  (template_sections)
    Content    = AI-generated text           (section_content)
    """
    doc = Document()

    # ── Cover page ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_heading("LOW LEVEL DESIGN", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_heading("Interface Architecture Document", level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph()

    # Metadata block
    meta_table = doc.add_table(rows=4, cols=2, style="Table Grid")
    meta_data  = [
        ("Interface ID",     interface_id or f"INT-{source[:3].upper()}-{datetime.now().strftime('%Y%m%d')}-001"),
        ("Source System",    source),
        ("Target System(s)", ", ".join(targets)),
        ("Generated On",     datetime.now().strftime("%d %B %Y  %H:%M")),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = k
        meta_table.rows[i].cells[1].text = v
        for run in meta_table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_page_break()

    # ── Write each template section ────────────────────────────────────────
    for section in template_sections:
        level = min(section["level"], 4)   # Word supports Heading 1-4 well
        title = section["title"]

        doc.add_heading(title, level=level)

        content = section_content.get(title, "")
        _write_content(doc, content)

    doc.save(output_path)
    return output_path

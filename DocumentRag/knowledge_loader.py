"""
knowledge_loader.py
Processes uploaded knowledge base files and indexes them into Cosmos DB.
Also extracts section structure from uploaded interface template .docx.
"""

import tempfile, os, traceback
from dotenv import load_dotenv
load_dotenv()

from pathlib import Path
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain.schema import Document as LCDocument  # only needed by disabled _load_md
from docx import Document
from cosmos_store import ensure_container_exists, upsert_many

LOADER_MAP     = {
    ".pdf":  PyPDFLoader,
    ".docx": Docx2txtLoader,
    ".txt":  lambda path: TextLoader(path, encoding="utf-8"),
}
# .md temporarily disabled — single source of truth is kb_cis.txt
SUPPORTED_TYPES = [".pdf", ".docx", ".txt"]  # , ".md"]


# def _load_md(path: str) -> list:
#     with open(path, "r", encoding="utf-8", errors="ignore") as f:
#         text = f.read()
#     return [LCDocument(page_content=text, metadata={"source": path})]


def _save_upload_to_temp(uploaded_file) -> str:
    suffix = Path(uploaded_file.name).suffix.lower()
    uploaded_file.seek(0)                    # reset Streamlit file pointer
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.write(uploaded_file.read())
    tmp.flush()
    tmp.close()                              # close before loader reads (Windows)
    return tmp.name


def build_vectorstore_from_uploads(uploaded_files: list, **kwargs) -> int:
    ensure_container_exists()
    splitter   = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
    all_chunks = []
    errors     = []

    for uf in uploaded_files:
        suffix = Path(uf.name).suffix.lower()

        if suffix not in SUPPORTED_TYPES:
            errors.append(f"{uf.name}: unsupported type '{suffix}'")
            continue

        tmp_path = _save_upload_to_temp(uf)
        try:
            # .md branch disabled — see _load_md commented out above
            # raw_docs = _load_md(tmp_path) if suffix == ".md" \
            #            else LOADER_MAP[suffix](tmp_path).load()
            raw_docs = LOADER_MAP[suffix](tmp_path).load()

            if not raw_docs:
                errors.append(f"{uf.name}: loader returned empty content")
                continue

            chunks = splitter.split_documents(raw_docs)
            if not chunks:
                errors.append(f"{uf.name}: no chunks after splitting")
                continue

            for chunk in chunks:
                all_chunks.append({
                    "text":         chunk.page_content,
                    "source_file":  uf.name,
                    "pattern_type": "uploaded",
                })
            print(f"  {uf.name} → {len(chunks)} chunks")

        except Exception as e:
            errors.append(f"{uf.name}: {str(e)}")
            traceback.print_exc()           # full error in terminal
        finally:
            try:
                os.unlink(tmp_path)
            except Exception:
                pass

    if errors:
        raise ValueError("\n".join(errors))

    if all_chunks:
        upsert_many(all_chunks)

    return len(all_chunks)


def extract_template_sections(uploaded_template) -> list[dict]:
    tmp_path = _save_upload_to_temp(uploaded_template)
    try:
        doc = Document(tmp_path)
        sections = []
        for para in doc.paragraphs:
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                try:
                    level = int(para.style.name.split()[-1])
                except ValueError:
                    level = 1
                if para.text.strip():
                    sections.append({"level": level, "title": para.text.strip()})
        return sections
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass